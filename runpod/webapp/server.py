"""Self-hosted transcribe-anything webapp.

Single-process FastAPI service that:
  1. Accepts an input URL via POST /api/jobs
  2. Optionally pre-fetches the audio through a residential-IP resolver host
     (configurable via env) so YouTube / Spotify / paywalled URLs work
  3. Submits the resulting audio URL to a RunPod Serverless transcribe-anything
     endpoint
  4. Polls RunPod and updates an in-memory job state
  5. Returns the full transcript + diarization on GET /api/jobs/{job_id}

Serves the static frontend at /. Intended to run behind a Tailscale-only
listener; no app-level auth.

Environment variables (loaded from os.environ; also auto-loads $HOME/.vgh.env):
  RUNPOD_API_KEY                required
  RUNPOD_ENDPOINT_ID            required
  HF_TOKEN                      optional — passed to RunPod for diarization
  TRANSCRIBE_RESOLVER_HOST      required if --resolve is used
                                (SSH target, e.g. user@host)
  TRANSCRIBE_RESOLVER_SCRIPT    required if --resolve is used
                                (absolute path to resolve_and_host.py on host)
  TRANSCRIBE_BIND_HOST          default: 127.0.0.1  (override to Tailscale IP)
  TRANSCRIBE_BIND_PORT          default: 8050
"""

from __future__ import annotations

import asyncio
import json
import os
import re
import subprocess
import sys
import time
import uuid
from pathlib import Path
from typing import Any, Optional

import httpx
from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel


def _load_vgh_env() -> None:
    """Best-effort loader for ~/.vgh.env (shell-style KEY=value)."""
    path = Path.home() / ".vgh.env"
    if not path.is_file():
        return
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        k, v = k.strip(), v.strip().strip('"').strip("'")
        if k and k not in os.environ:
            os.environ[k] = v


_load_vgh_env()


def _env(name: str, required: bool = False, default: str = "") -> str:
    val = os.environ.get(name, default)
    if required and not val:
        sys.stderr.write(f"FATAL: required env var {name} not set\n")
        sys.exit(2)
    return val


RUNPOD_API_KEY = _env("RUNPOD_API_KEY", required=True)
RUNPOD_ENDPOINT_ID = _env("RUNPOD_ENDPOINT_ID", required=True)
HF_TOKEN = _env("HF_TOKEN")
RESOLVER_HOST = _env("TRANSCRIBE_RESOLVER_HOST")
RESOLVER_SCRIPT = _env("TRANSCRIBE_RESOLVER_SCRIPT")
BIND_HOST = _env("TRANSCRIBE_BIND_HOST", default="127.0.0.1")
BIND_PORT = int(_env("TRANSCRIBE_BIND_PORT", default="8050"))

# Obsidian note creation — if both are set, write a markdown note per
# completed job into the vault on a remote host.
OBSIDIAN_HOST = _env("OBSIDIAN_HOST")                    # SSH target, e.g. user@host
OBSIDIAN_VAULT_DIR = _env("OBSIDIAN_VAULT_DIR")          # absolute path on host, e.g. /home/.../VKJOS2025/00_Raw/Transcripts

RUNPOD_BASE = f"https://api.runpod.ai/v2/{RUNPOD_ENDPOINT_ID}"
RUNPOD_HEADERS = {"Authorization": f"Bearer {RUNPOD_API_KEY}"}


# --- Job state -------------------------------------------------------------

# JSON-file-backed job store. In-memory dict mirrored to JOBS_FILE on every
# state change so a service restart doesn't lose in-flight work. Single-file,
# best-effort fsync; swap for SQLite or Redis if concurrent writes become
# a concern.
JOBS_FILE = Path(_env("TRANSCRIBE_JOBS_FILE", default="/var/lib/transcribe-webapp/jobs.json"))
JOBS: dict[str, dict[str, Any]] = {}


def _load_jobs_from_disk() -> None:
    """Best-effort restore at startup."""
    if not JOBS_FILE.is_file():
        return
    try:
        data = json.loads(JOBS_FILE.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            JOBS.update(data)
            sys.stderr.write(f"[startup] restored {len(JOBS)} job(s) from {JOBS_FILE}\n")
    except Exception as e:
        sys.stderr.write(f"[startup] could not restore JOBS from {JOBS_FILE}: {e}\n")


def _persist_jobs() -> None:
    """Write JOBS to disk. Atomic via write-temp-then-rename."""
    try:
        JOBS_FILE.parent.mkdir(parents=True, exist_ok=True)
        tmp = JOBS_FILE.with_suffix(JOBS_FILE.suffix + ".tmp")
        tmp.write_text(json.dumps(JOBS, default=str), encoding="utf-8")
        tmp.replace(JOBS_FILE)
    except Exception as e:
        sys.stderr.write(f"[persist] failed: {e}\n")


_load_jobs_from_disk()


def _mask(s: Any) -> Any:
    """Strip hf_ / rpa_ tokens from any string before exposing in API responses."""
    if isinstance(s, str):
        s = re.sub(r"hf_[A-Za-z0-9]{8,}", "hf_<REDACTED>", s)
        s = re.sub(r"rpa_[A-Za-z0-9]{8,}", "rpa_<REDACTED>", s)
    elif isinstance(s, dict):
        return {k: _mask(v) for k, v in s.items()}
    elif isinstance(s, list):
        return [_mask(x) for x in s]
    return s


# --- App -------------------------------------------------------------------

app = FastAPI(title="transcribe-anything", version="0.1.0")


class SubmitRequest(BaseModel):
    url: str
    use_resolver: bool = False
    model: str = "large-v3"


@app.post("/api/jobs")
async def submit(req: SubmitRequest) -> dict[str, str]:
    if not req.url.strip():
        raise HTTPException(400, "url is empty")
    if req.use_resolver and (not RESOLVER_HOST or not RESOLVER_SCRIPT):
        raise HTTPException(
            400,
            "use_resolver=true but TRANSCRIBE_RESOLVER_HOST/SCRIPT not configured",
        )
    job_id = uuid.uuid4().hex[:12]
    JOBS[job_id] = {
        "job_id": job_id,
        "status": "submitted",
        "input_url": req.url,
        "use_resolver": req.use_resolver,
        "model": req.model,
        "created_at": time.time(),
        "progress_message": "queued",
        "runpod_id": None,
        "result": None,
        "error": None,
    }
    _persist_jobs()
    asyncio.create_task(_process_job(job_id))
    return {"job_id": job_id, "status": "submitted"}


@app.get("/api/jobs/{job_id}")
async def get_job(job_id: str) -> dict[str, Any]:
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(404, "job not found")
    return _mask(job)


@app.get("/api/jobs/{job_id}/markdown")
async def get_job_markdown(job_id: str):
    """Download the completed transcript as an Obsidian-style .md file."""
    from fastapi.responses import Response
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(404, "job not found")
    if job.get("status") != "completed":
        raise HTTPException(409, f"job not completed yet (status: {job.get('status')})")
    filename, body = _build_obsidian_note(job)
    return Response(
        content=body,
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/jobs/{job_id}/docx")
async def get_job_docx(job_id: str):
    """Download the completed transcript as a Word .docx file."""
    from fastapi.responses import Response
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(404, "job not found")
    if job.get("status") != "completed":
        raise HTTPException(409, f"job not completed yet (status: {job.get('status')})")
    filename, buf = _build_docx(job)
    return Response(
        content=buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/api/health")
async def health() -> dict[str, Any]:
    return {
        "ok": True,
        "endpoint_id": RUNPOD_ENDPOINT_ID,
        "resolver_configured": bool(RESOLVER_HOST and RESOLVER_SCRIPT),
        "hf_token_configured": bool(HF_TOKEN),
        "obsidian_configured": bool(OBSIDIAN_HOST and OBSIDIAN_VAULT_DIR),
        "jobs_in_memory": len(JOBS),
    }


# --- Pipeline --------------------------------------------------------------

async def _run_resolver(input_url: str) -> str:
    """SSH to RESOLVER_HOST and run RESOLVER_SCRIPT, returning the resolved URL.

    OpenSSH joins all remote-command args into a single string for the remote
    shell to evaluate, so any shell-special chars in input_url (?, *, $, etc.)
    must be quoted before sending. YouTube URLs in particular contain `?v=...`
    which zsh on the resolver host interprets as a glob and fails with
    'no matches found'. shlex.quote() makes the URL literal on the remote side.
    """
    import shlex
    remote_cmd = (
        "python3 "
        + shlex.quote(RESOLVER_SCRIPT)
        + " "
        + shlex.quote(input_url)
    )
    cmd = [
        "ssh", "-o", "ConnectTimeout=10",
        RESOLVER_HOST, remote_cmd,
    ]
    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, stderr = await proc.communicate()
    if proc.returncode != 0:
        raise RuntimeError(
            f"resolver exited {proc.returncode}: {stderr.decode(errors='replace')[:500]}"
        )
    last_line = stdout.decode(errors="replace").strip().splitlines()[-1] if stdout.strip() else ""
    if not last_line.startswith("http"):
        raise RuntimeError(f"resolver did not return a URL; got: {last_line!r}")
    return last_line


async def _runpod_submit(audio_url: str, model: str) -> str:
    payload = {
        "input": {
            "url_or_file": audio_url,
            "device": "insane",
            "model": model,
            "task": "transcribe",
            "flash": False,
            "batch_size": 8,
            "hf_token": HF_TOKEN or None,
        }
    }
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.post(f"{RUNPOD_BASE}/run", headers=RUNPOD_HEADERS, json=payload)
        r.raise_for_status()
        return r.json()["id"]


async def _runpod_status(runpod_id: str) -> dict[str, Any]:
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(f"{RUNPOD_BASE}/status/{runpod_id}", headers=RUNPOD_HEADERS)
        r.raise_for_status()
        return r.json()


def _slugify(s: str, max_len: int = 60) -> str:
    """Conservative filename slug — keeps alnum + spaces, collapses runs."""
    s = re.sub(r"[^\w\s-]+", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s[:max_len].strip()


def _build_obsidian_note(job: dict[str, Any]) -> tuple[str, str]:
    """Build (filename, markdown_body) for a completed transcription job.

    Filename follows the vault convention: 'YYYY-MM-DD <title>.md'.
    """
    from datetime import datetime, timezone
    result = job.get("result") or {}
    speaker_json = result.get("speaker_json") or []
    text = result.get("text") or ""

    # Title derivation: first 60 chars of transcript text, or the URL host.
    if text:
        title = _slugify(text.split(".")[0]) or _slugify(text[:60])
    else:
        title = _slugify(job.get("input_url", "transcript").split("/")[-1] or "transcript")
    if not title:
        title = "transcript"

    date_str = datetime.fromtimestamp(job.get("completed_at", time.time()), tz=timezone.utc).strftime("%Y-%m-%d")
    filename = f"{date_str} {title}.md"

    # Frontmatter
    speakers = sorted({seg.get("speaker") for seg in speaker_json if isinstance(seg, dict) and seg.get("speaker")})
    duration = 0.0
    if speaker_json:
        last = speaker_json[-1]
        if isinstance(last, dict) and isinstance(last.get("timestamp"), list) and len(last["timestamp"]) >= 2:
            duration = last["timestamp"][1] or 0.0
    duration_str = f"{int(duration // 60)}m {int(duration % 60)}s" if duration else ""

    lines: list[str] = []
    lines.append("---")
    lines.append("type: transcript")
    lines.append(f"created: {date_str}")
    lines.append(f"source_url: {job.get('input_url', '')}")
    if job.get("resolved_url"):
        lines.append(f"resolved_url: {job['resolved_url']}")
    lines.append(f"runpod_job_id: {job.get('runpod_id', '')}")
    if speakers:
        lines.append(f"speakers: {len(speakers)}")
        lines.append(f"speaker_labels: [{', '.join(speakers)}]")
    if duration_str:
        lines.append(f"duration: {duration_str}")
    lines.append("tags: [transcript, podcast]")
    lines.append("---")
    lines.append("")
    lines.append(f"# {title}")
    lines.append("")
    lines.append(f"**Source**: {job.get('input_url', '')}")
    lines.append("")

    # Speaker-grouped transcript
    if speaker_json:
        lines.append("## Transcript")
        lines.append("")
        for seg in speaker_json:
            if not isinstance(seg, dict):
                continue
            spk = seg.get("speaker", "SPEAKER_??")
            ts = seg.get("timestamp") or [0, 0]
            mm = int((ts[0] or 0) // 60)
            ss = int((ts[0] or 0) % 60)
            ts_str = f"`[{mm:02d}:{ss:02d}]`"
            seg_text = (seg.get("text") or "").strip()
            lines.append(f"**{spk}** {ts_str}")
            lines.append("")
            lines.append(seg_text)
            lines.append("")
    else:
        lines.append("## Transcript")
        lines.append("")
        lines.append(text or "*(no transcript text)*")
        lines.append("")

    return filename, "\n".join(lines)


def _build_docx(job: dict[str, Any]) -> tuple[str, bytes]:
    """Build (filename, raw_docx_bytes) for a completed job.

    Real .docx (Word XML), not the HTML-pretending-to-be-doc trick. Uses
    python-docx for proper headings, bold speaker labels, and timestamps.
    """
    from datetime import datetime, timezone
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor

    result = job.get("result") or {}
    speaker_json = result.get("speaker_json") or []
    text = result.get("text") or ""

    # Mirror the .md title derivation so the two files line up
    if text:
        title = _slugify(text.split(".")[0]) or _slugify(text[:60])
    else:
        title = _slugify(job.get("input_url", "transcript").split("/")[-1] or "transcript")
    if not title:
        title = "transcript"
    date_str = datetime.fromtimestamp(
        job.get("completed_at", time.time()), tz=timezone.utc
    ).strftime("%Y-%m-%d")
    filename = f"{date_str} {title}.docx"

    doc = Document()
    # Set a readable default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    meta = doc.add_paragraph()
    meta.add_run("Source: ").bold = True
    meta.add_run(job.get("input_url", ""))
    if job.get("resolved_url"):
        meta.add_run("\nResolved: ").bold = True
        meta.add_run(job["resolved_url"])

    speakers = sorted({seg.get("speaker") for seg in speaker_json if isinstance(seg, dict) and seg.get("speaker")})
    duration = 0.0
    if speaker_json:
        last = speaker_json[-1]
        if isinstance(last, dict) and isinstance(last.get("timestamp"), list) and len(last["timestamp"]) >= 2:
            duration = last["timestamp"][1] or 0.0

    stats = doc.add_paragraph()
    stats.add_run("Stats: ").bold = True
    stats.add_run(f"{len(speakers)} speakers, {len(speaker_json)} turns")
    if duration:
        stats.add_run(f", {int(duration // 60)}m {int(duration % 60)}s duration")

    doc.add_heading("Transcript", level=1)

    if speaker_json:
        # Color cycle that loosely matches the webapp UI palette
        speaker_colors = [
            RGBColor(0xCC, 0xA0, 0x00),  # darker yellow for contrast on white
            RGBColor(0xC0, 0x4A, 0x9B),  # blush pink
            RGBColor(0x4A, 0x6A, 0xD0),  # sky blue
            RGBColor(0x1F, 0x80, 0x76),  # grass green
            RGBColor(0xB0, 0x60, 0x60),  # rust
            RGBColor(0x80, 0x60, 0xB0),  # purple
        ]
        speaker_color_map: dict[str, RGBColor] = {}
        for seg in speaker_json:
            if not isinstance(seg, dict):
                continue
            spk = seg.get("speaker", "SPEAKER_??")
            if spk not in speaker_color_map:
                speaker_color_map[spk] = speaker_colors[len(speaker_color_map) % len(speaker_colors)]
            ts = seg.get("timestamp") or [0, 0]
            mm, ss = int((ts[0] or 0) // 60), int((ts[0] or 0) % 60)
            seg_text = (seg.get("text") or "").strip()

            p = doc.add_paragraph()
            spk_run = p.add_run(spk)
            spk_run.bold = True
            spk_run.font.color.rgb = speaker_color_map[spk]
            ts_run = p.add_run(f"  [{mm:02d}:{ss:02d}]  ")
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            p.add_run(seg_text)
    else:
        doc.add_paragraph(text or "(no transcript text)")

    buf = BytesIO()
    doc.save(buf)
    return filename, buf.getvalue()


async def _write_obsidian_note(job: dict[str, Any]) -> Optional[str]:
    """SSH to OBSIDIAN_HOST and write the note into OBSIDIAN_VAULT_DIR.

    Returns the written path on success, None if Obsidian export isn't
    configured. Raises on actual failure (caller decides what to do).
    """
    if not OBSIDIAN_HOST or not OBSIDIAN_VAULT_DIR:
        return None
    filename, body = _build_obsidian_note(job)
    remote_path = f"{OBSIDIAN_VAULT_DIR.rstrip('/')}/{filename}"
    # Pipe the body to ssh which writes it to the file. tee is robust
    # (creates the dir if missing via prior mkdir).
    cmd = [
        "ssh", "-o", "ConnectTimeout=10", OBSIDIAN_HOST,
        f"mkdir -p {OBSIDIAN_VAULT_DIR!r} && cat > {remote_path!r}",
    ]
    proc = await asyncio.create_subprocess_exec(
        *cmd,
        stdin=asyncio.subprocess.PIPE,
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
    )
    stdout, stderr = await proc.communicate(input=body.encode("utf-8"))
    if proc.returncode != 0:
        raise RuntimeError(
            f"obsidian write exited {proc.returncode}: {stderr.decode(errors='replace')[:300]}"
        )
    return remote_path


async def _process_job(job_id: str) -> None:
    job = JOBS[job_id]
    try:
        audio_url = job["input_url"]
        if job["use_resolver"]:
            job["status"] = "resolving"
            job["progress_message"] = "fetching audio through resolver host"
            _persist_jobs()
            audio_url = await _run_resolver(audio_url)
            job["resolved_url"] = audio_url

        job["status"] = "submitting"
        job["progress_message"] = "submitting to RunPod"
        _persist_jobs()
        runpod_id = await _runpod_submit(audio_url, job["model"])
        job["runpod_id"] = runpod_id

        job["status"] = "transcribing"
        job["progress_message"] = "RunPod is transcribing (this can take several minutes for long audio)"
        _persist_jobs()

        # Poll until terminal
        for _ in range(120):  # 120 * 15s = 30 min max
            await asyncio.sleep(15)
            d = await _runpod_status(runpod_id)
            s = d.get("status")
            if s == "COMPLETED":
                job["status"] = "completed"
                job["progress_message"] = "done"
                job["result"] = d.get("output") or {}
                job["completed_at"] = time.time()
                # Write to Obsidian vault if configured. Best-effort: a
                # failure here doesn't change the job result.
                try:
                    path = await _write_obsidian_note(job)
                    if path:
                        job["obsidian_path"] = path
                        job["progress_message"] = f"saved to {path}"
                except Exception as obs_err:
                    job["obsidian_error"] = f"{type(obs_err).__name__}: {str(obs_err)[:200]}"
                _persist_jobs()
                return
            if s in ("FAILED", "CANCELLED", "TIMED_OUT"):
                job["status"] = "failed"
                job["error"] = (d.get("error") or "")[:1500]
                job["progress_message"] = f"RunPod returned {s}"
                _persist_jobs()
                return
            # Update progress on intermediate states
            prev_msg = job.get("progress_message")
            if s == "IN_PROGRESS":
                job["progress_message"] = "RunPod worker is processing"
            elif s == "IN_QUEUE":
                job["progress_message"] = "queued on RunPod"
            if job.get("progress_message") != prev_msg:
                _persist_jobs()
        # Loop exited without terminal status — give up
        job["status"] = "failed"
        job["error"] = "exceeded 30 min poll window"
        _persist_jobs()
    except Exception as e:
        job["status"] = "failed"
        job["error"] = f"{type(e).__name__}: {str(e)[:500]}"
        _persist_jobs()


# --- Static frontend -------------------------------------------------------

_STATIC_DIR = Path(__file__).parent / "static"
if _STATIC_DIR.is_dir():
    app.mount("/", StaticFiles(directory=str(_STATIC_DIR), html=True), name="static")


# --- CLI entrypoint --------------------------------------------------------

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host=BIND_HOST, port=BIND_PORT, log_level="info")
